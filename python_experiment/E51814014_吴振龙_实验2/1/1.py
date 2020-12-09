import random
import os
import docx

operators = ['+', '-', '*', '/']

def generateExpression(operatorNum):
    ''' 随机生成一个operatorNum个运算符的四则运算式 '''
    # 表达式字符串
    expressionString = ""
    numbersList = []            # 生成的所有数字
    leftBracket = 0             # 一个表达式中有一个未匹配的左括号的数量
    for i in range(operatorNum + 1):
        # 前operatorNum次循环中，每次生成一个运算数和一个运算符，最后一次循环只生成一个运算数
        leftBracketThisTurn =  random.randint(0, 1)
        if leftBracketThisTurn and i != operatorNum:
            # 随机到1且本轮不是生成最后一个运算符时，生成一个左括号
            expressionString += '('
            leftBracket += 1
        # 生成数字
        if expressionString and expressionString[-1] == '/':
            # 如果最后一个运算符是除，不生成1
            numbersList.append(random.randint(2, 99))
        else:
            numbersList.append(random.randint(1, 99))
        expressionString += str(numbersList[-1])
        if i == operatorNum:
            # 如果已经生成了足够多的运算符，在此结束，不再生成运算符
            # 在运算式末尾补全右括号
            for i in range(leftBracket):
                expressionString += ')'
            break
        if not leftBracketThisTurn and leftBracket>0 and random.randint(0, 1):
            # 如果本轮没有生成过左括号且此时有未匹配过的左括号随机到1，生成一个右括号
            expressionString += ')'
            leftBracket -= 1
        # 随机生成一个运算符
        if numbersList[-1] == 1:
            # 如果最后一个数字是1，运算符不生成乘号
            expressionString += ['+', '-', '/'][random.randint(0, 2)]
        else:
            expressionString += operators[random.randint(0, 3)]

    try:
        eval(expressionString)
    except ZeroDivisionError:
        # 表达式中出现0除，返回空
        return None, None
    answer = eval(expressionString)
    # 答案保留两位小数
    answer = "{:.2f}".format(answer)
    return expressionString, answer
# 获得题目和答案文件路径
questionsPath = os.path.dirname(__file__) + "/题目.docx"
answersPath = os.path.dirname(__file__) + "/参考答案.docx"
questionsNum = 50
expressions = []
answers = []
for i in range(questionsNum):
    expressionLen = random.randint(3, 9)
    expression, answer = generateExpression(expressionLen)
    # 在每个元素字符串后面加上换行，方便写入文件
    expressions.append(expression + '\n')
    answers.append(answer + '\n')
# 创建文档对象
questionsDocument = docx.Document()
answersDocument = docx.Document()
for i in range(len(expressions)):
    # 以有序列表形式写入文档
    questionsDocument.add_paragraph(expressions[i], style = 'List Number')
    answersDocument.add_paragraph(answers[i], style = 'List Number')
questionsDocument.save(questionsPath)
answersDocument.save(answersPath)